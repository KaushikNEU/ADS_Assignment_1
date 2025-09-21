# ------------------------------------------------------------
# Botspeak | DIRECT (Interactive) — Format 1 + Quiz
# - Uses .env (OPENAI_API_KEY) OR Streamlit secrets for OpenAI
# - Reads "Botspeak - Assignment 1.docx" by default OR via uploader
# - Story -> Prompt Builder -> LLM Generate & Acceptance Tests -> Quiz
# ------------------------------------------------------------

import os, re, io
from typing import List, Tuple, Dict

import numpy as np
import streamlit as st

# ---- .env (OpenAI key) ----
try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")

# ---- OpenAI client (optional but recommended) ----
USE_OPENAI = False
client = None
MODEL_DEFAULT = "gpt-4o-mini"  # change if you prefer another available model
SYSTEM_HINT = "You are a careful assistant who follows prompt specs exactly and adheres to safety rules."

try:
    from openai import OpenAI
    if OPENAI_API_KEY:
        client = OpenAI(api_key=OPENAI_API_KEY)
        USE_OPENAI = True
except Exception:
    USE_OPENAI = False
    client = None

# ---- DOCX parsing + retrieval ----
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

DOC_PATH_DEFAULT = "Botspeak - Assignment 1.docx"

@st.cache_data(show_spinner=False)
def load_docx_paragraphs_from_path(path: str) -> List[str]:
    doc = docx.Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # remove near-duplicates
    keep, seen = [], set()
    for p in paras:
        k = " ".join(p.lower().split())
        if k not in seen:
            seen.add(k)
            keep.append(p)
    return keep

def load_docx_paragraphs_from_bytes(data: bytes) -> List[str]:
    d = docx.Document(io.BytesIO(data))
    paras = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    keep, seen = [], set()
    for p in paras:
        k = " ".join(p.lower().split())
        if k not in seen:
            seen.add(k)
            keep.append(p)
    return keep

@st.cache_resource(show_spinner=False)
def build_index(paras: List[str]):
    vec = TfidfVectorizer(ngram_range=(1,2), min_df=1)
    X = vec.fit_transform(paras)
    return vec, X

def retrieve_explainer(query: str, vec, X, paras: List[str], k: int = 2) -> List[str]:
    if not paras or vec is None or X is None:
        return []
    q = vec.transform([query])
    sims = cosine_similarity(q, X).ravel()
    top = sims.argsort()[::-1][:k]
    return [paras[i] for i in top]

# ---- Prompt Spec + Tests ----
def assemble_good_prompt(role, goal, context, constraints, schema, examples_good, example_bad, tests):
    blocks = [
        f"ROLE:\n{role}",
        f"GOAL:\n{goal}",
        f"CONTEXT / INPUTS:\n{context}",
        f"CONSTRAINTS & SAFETY RULES:\n{constraints}",
        f"OUTPUT SCHEMA:\n{schema}",
        "FEW-SHOT EXAMPLES (GOOD):\n" + examples_good,
        "COUNTER-EXAMPLE (BAD):\n" + example_bad,
        "ACCEPTANCE TESTS:\n" + tests,
    ]
    return "\n\n---\n\n".join(blocks)

def assemble_vague_prompt(goal):
    return f"Please do this:\n{goal}\nThanks."

def run_acceptance_tests(outputs: List[str], max_words=50, banned=None, must_be_numbered=True) -> Tuple[List[Dict[str,bool]], List[bool]]:
    banned = [b.strip().lower() for b in (banned or []) if b.strip()]
    results = []
    for out in outputs:
        checks = {}
        checks["≤ word_limit"] = (len(out.split()) <= max_words)
        lo = out.lower()
        checks["no_banned_terms"] = not any(bt in lo for bt in banned)
        checks["single_line"] = ("\n" not in out.strip())
        if must_be_numbered:
            checks["numbered_item"] = re.match(r"^\s*\d+[\).\-:\s]", out.strip()) is not None
        else:
            checks["numbered_item"] = True
        results.append(checks)
    passes = [all(d.values()) for d in results]
    return results, passes

def llm_generate_with_openai(prompt_spec: str, n: int, model: str, temperature: float) -> List[str]:
    user_msg = prompt_spec + f"\n\nGenerate exactly {n} items only."
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": SYSTEM_HINT},
            {"role": "user", "content": user_msg},
        ],
        temperature=float(temperature),
    )
    text = resp.choices[0].message.content.strip()
    items = []
    for line in text.splitlines():
        m = re.match(r"^\s*(\d+)[\).\-:\s](.*)$", line.strip())
        if m:
            items.append(f"{m.group(1)}. {m.group(2).strip()}")
        elif line.strip():
            items.append(line.strip())
    return items[:n]

def llm_stub(n: int) -> List[str]:
    return [f"{i}. Master time, defy logic—ChronoShift deals out brainy twists in under 50 words." for i in range(1, n+1)]

# ---- High-contrast chips & cards (theme-aware) ----
BADGE_CSS = """
<style>
:root{
  --chip-fg:#0f172a; --chip-bg:#e5e8f0; --chip-bd:#cbd5e1;
  --chip-good-bg:#e8fff4; --chip-good-bd:#22c55e;
  --chip-warn-bg:#fff7ed; --chip-warn-bd:#f59e0b;
  --chip-dang-bg:#fff1f2; --chip-dang-bd:#ef4444;
  --card-bg:#f8fafc; --card-bd:#e5e7eb;
}
@media (prefers-color-scheme: dark){
  :root{
    --chip-fg:#e5e7eb; --chip-bg:#111827; --chip-bd:#374151;
    --chip-good-bg:#064e3b; --chip-good-bd:#34d399;
    --chip-warn-bg:#7c2d12; --chip-warn-bd:#fbbf24;
    --chip-dang-bg:#7f1d1d; --chip-dang-bd:#f87171;
    --card-bg:#0b1220; --card-bd:#1f2937;
  }
}
.chip{display:inline-flex;align-items:center;gap:6px;padding:6px 10px;border-radius:999px;
  font-size:12px;line-height:1;font-weight:600;color:var(--chip-fg);
  background:var(--chip-bg);border:1px solid var(--chip-bd);}
.chip.good{background:var(--chip-good-bg);border-color:var(--chip-good-bd);}
.chip.warn{background:var(--chip-warn-bg);border-color:var(--chip-warn-bd);}
.chip.dang{background:var(--chip-dang-bg);border-color:var(--chip-dang-bd);}
.card{border:1px solid var(--card-bd);padding:12px;border-radius:12px;background:var(--card-bg);}
.status{display:flex;justify-content:flex-end;margin-top:-18px;}
</style>
"""

def api_status_badge():
    _, right = st.columns([1,0.33])
    with right:
        if USE_OPENAI:
            st.markdown("<div class='status'><span class='chip good'>🔐 OpenAI connected</span></div>", unsafe_allow_html=True)
        else:
            st.markdown("<div class='status'><span class='chip dang'>🔒 OpenAI key missing</span></div>", unsafe_allow_html=True)

# ---- App ----
st.set_page_config(page_title="Botspeak | DIRECT (Interactive)", layout="wide")
st.markdown(BADGE_CSS, unsafe_allow_html=True)

st.title("DIRECT the AI — Interactive Walkthrough")
st.caption("Botspeak Loop · Step 3 · Turn goals into a Prompt Spec with roles, schema, constraints, examples, and acceptance tests.")
api_status_badge()

# ---- DOCX load: default or via uploader (“Browse files”) ----
# The uploader below renders a "Browse files" button in Streamlit’s UI.
paras = []
doc_loaded = False
vec = X = None

default_exists = os.path.exists(DOC_PATH_DEFAULT)
up = st.file_uploader("(Optional) Upload your Botspeak DOCX to power quiz explanations", type=["docx"], accept_multiple_files=False)

if up is not None:
    try:
        paras = load_docx_paragraphs_from_bytes(up.read())
        vec, X = build_index(paras)
        doc_loaded = True
        st.success("Knowledge doc loaded from upload ✅")
    except Exception as e:
        st.error(f"Could not read the uploaded DOCX: {e}")
elif default_exists:
    try:
        paras = load_docx_paragraphs_from_path(DOC_PATH_DEFAULT)
        vec, X = build_index(paras)
        doc_loaded = True
        st.success(f"Knowledge doc loaded: {DOC_PATH_DEFAULT} ✅")
    except Exception as e:
        st.warning(f"Could not load '{DOC_PATH_DEFAULT}'. You can upload the DOCX above. ({e})")
else:
    st.info("Tip: Place your DOCX next to this file as "
            f"'{DOC_PATH_DEFAULT}' or upload it above for grounded quiz explanations.")

tabs = st.tabs(["Story", "Prompt Builder", "Generate & Test", "Quiz"])

# =============== STORY TAB ===============
with tabs[0]:
    st.subheader("What is “Direct” and why it matters")
    left, right = st.columns([1.25, 1])
    with left:
        st.markdown("""
- **Definition**: Convert tasks into a **Prompt Spec** the model can execute (role, inputs, constraints, schema, examples, **acceptance tests**).
- **Why it matters**:  
  <span class='chip good'>⚙️ Reliability</span>
  <span class='chip good'>🛡️ Safety</span>
  <span class='chip good'>🎯 Usefulness</span>
- **Bridge**: Turns strategy (Define/Delegate) into consistent execution, and sets criteria for Diagnose & Document.
""", unsafe_allow_html=True)
        st.info("Story setup: JK Games is launching **ChronoShift** and needs hundreds of on-brand ad descriptions for A/B tests.")
    with right:
        st.markdown("**Choose your hat**")
        role_choice = st.radio("You are…", ["Marketing Lead", "Creative Director", "AI Engineer"], label_visibility="collapsed")

    if role_choice == "Marketing Lead":
        st.success("Focus: CTR, volume, time-to-first-draft. Guard against off-brand or too-long outputs.")
    elif role_choice == "Creative Director":
        st.success("Focus: brand voice, banned terms, platform policies, ethics. Guardrails & counter-examples matter.")
    else:
        st.success("Focus: deterministic schema, test hooks, automation fit, and failure modes for Diagnose step.")

# =============== PROMPT BUILDER TAB ===============
with tabs[1]:
    st.subheader("Build your Prompt Spec")
    with st.form("spec_form"):
        c1, c2 = st.columns(2)
        with c1:
            role_txt = st.text_input("Role", "Senior marketing copywriter for a mobile game studio.")
            goal_txt = st.text_area("Goal", "Generate 50 unique ad descriptions (25–50 words) for ChronoShift.")
            context_txt = st.text_area("Context / Inputs",
                "Game: ChronoShift. Themes: time-bending, logic, puzzle-solving. Audience: brain-teaser fans. Platforms: Google Ads, Facebook.")
            schema_txt = st.text_area("Output Schema",
                "Numbered list. Each item is a single ad description (one line) without line breaks.")
        with c2:
            constraints_txt = st.text_area("Constraints & Safety Rules",
                "≤ 50 words per description. No claims about IQ/memory. Banned terms: addictive, life-changing, brain training.")
            examples_good = st.text_area("Two good examples",
                "- Master time, defy logic. ChronoShift challenges the way you think with mind-bending puzzles.\n- A new era of puzzle games. Shift time, crack the pattern, and outsmart the clock.")
            example_bad = st.text_area("One counter-example",
                "WARNING: Guaranteed IQ boost! The most awesome puzzle game ever made—buy now!")
            tests_txt = st.text_area("Acceptance tests",
                "- Each line ≤ 50 words\n- No banned terms (addictive, life-changing, brain training)\n- No features not in context\n- Must be a numbered list")

        cols = st.columns(3)
        with cols[0]:
            model_name = st.text_input("OpenAI model", MODEL_DEFAULT)
        with cols[1]:
            gen_preview_n = st.slider("Preview items to generate", 5, 50, 10, step=5)
        with cols[2]:
            temperature = st.slider("Creativity (temperature)", 0.0, 1.0, 0.7, 0.1)

        submitted = st.form_submit_button("Score & Save")
        if submitted:
            spec_fields = [role_txt, goal_txt, context_txt, constraints_txt, schema_txt, examples_good, example_bad, tests_txt]
            score = sum(bool(x.strip()) for x in spec_fields)
            st.session_state["spec"] = {
                "role": role_txt, "goal": goal_txt, "context": context_txt,
                "constraints": constraints_txt, "schema": schema_txt,
                "good": examples_good, "bad": example_bad, "tests": tests_txt,
                "model": model_name, "n": gen_preview_n, "temp": temperature
            }
            st.success(f"Prompt Spec completeness: **{score}/8**")

            checks = {
                "Output schema present": bool(schema_txt.strip()),
                "Constraints & safety rules present": bool(constraints_txt.strip()),
                "Examples + counter-example": bool(examples_good.strip() and example_bad.strip()),
                "Acceptance tests present": bool(tests_txt.strip())
            }
            bcols = st.columns(4)
            for (k, v), c in zip(checks.items(), bcols):
                c.markdown(f"<span class='chip {'good' if v else 'dang'}'>{'✅' if v else '✖'} {k}</span>", unsafe_allow_html=True)

# =============== GENERATE & TEST TAB ===============
with tabs[2]:
    st.subheader("LLM Generation + Acceptance Tests")
    if "spec" not in st.session_state:
        st.warning("Build and save a Prompt Spec first (previous tab).")
    else:
        spec = st.session_state["spec"]
        good_prompt = assemble_good_prompt(
            spec["role"], spec["goal"], spec["context"],
            spec["constraints"], spec["schema"], spec["good"], spec["bad"], spec["tests"]
        )
        vague_prompt = assemble_vague_prompt(spec["goal"])

        c1, c2 = st.columns(2)
        with c1:
            st.caption("Good Prompt")
            st.code(good_prompt, language="markdown")
        with c2:
            st.caption("Vague Prompt (for contrast)")
            st.code(vague_prompt, language="markdown")

        gen_n = st.slider("How many items to generate?", 5, 100, spec.get("n", 10), step=5)
        go = st.button("Generate with LLM")
        if go:
            with st.spinner("Generating…"):
                if USE_OPENAI and client is not None:
                    try:
                        outputs = llm_generate_with_openai(
                            prompt_spec=good_prompt,
                            n=gen_n,
                            model=spec.get("model", MODEL_DEFAULT),
                            temperature=float(spec.get("temp", 0.7)),
                        )
                    except Exception as e:
                        st.error(f"OpenAI error: {e}")
                        outputs = llm_stub(gen_n)
                else:
                    outputs = llm_stub(gen_n)

            st.toast(f"Generated {len(outputs)} items.", icon="✅")
            st.write("### Output (first 12)")
            st.markdown("<div class='card'>", unsafe_allow_html=True)
            st.write("\n".join(outputs[:12]))
            st.markdown("</div>", unsafe_allow_html=True)

            # parse banned terms from constraints
            banned_terms = []
            m = re.search(r"banned terms\s*:\s*(.+)", spec["constraints"], flags=re.I)
            if m:
                banned_terms = [x.strip().lower().strip(".,;") for x in re.split(r"[,\|/]", m.group(1)) if x.strip()]

            results, passes = run_acceptance_tests(outputs, max_words=50, banned=banned_terms, must_be_numbered=True)
            passed = sum(passes)
            st.write(f"### Acceptance tests: {passed}/{len(outputs)} passed")

            for i, (out, ok, checkdict) in enumerate(zip(outputs, passes, results), 1):
                badge = "<span class='chip good'>PASS</span>" if ok else "<span class='chip dang'>FAIL</span>"
                st.markdown(f"{badge} **{i}.** {out}", unsafe_allow_html=True)
                st.caption(" · ".join([f"{'✅' if v else '❌'} {k}" for k, v in checkdict.items()]))

# =============== QUIZ TAB ===============
with tabs[3]:
    st.subheader("Knowledge Check (auto-graded, with explanations)")
    if not doc_loaded:
        st.warning("Explanations will be stronger when your DOCX is loaded (upload above or place it next to the app).")

    QUIZ = [
        {"type":"mcq",
         "q":"Which element makes the DIRECT step falsifiable?",
         "opts":["Tone guide","Output schema","Acceptance tests","Examples"],
         "a":2, "hint":"acceptance tests falsifiability reliability"},
        {"type":"scenario",
         "q":"Your ChronoShift prompt lacks banned terms & safety rules. Which risk is most likely to show first?",
         "opts":["Poor CTR","Legal/ethical rejection","Longer copy","Token limits"],
         "a":1, "hint":"guardrails refusal safety rules platform policies"},
        {"type":"application",
         "q":"Write one constraint that reduces hallucinations for ChronoShift.",
         "rubric":["Reference provided context only","Prohibit unstated features"],
         "hint":"constraints reduce hallucination context-bound falsifiability"},
        {"type":"mcq",
         "q":"Which pair best represents the DIRECT step’s foundations?",
         "opts":["Hume & Kant","Popper & Descartes","Bentham & Mill","Kuhn & Feyerabend"],
         "a":1, "hint":"Popper falsifiability Descartes methodic doubt structure"},
        {"type":"scenario",
         "q":"A bank asks for credit line increases with no schema/safety. Which failure appears first?",
         "opts":["Too short","Biased/inconsistent recs","Slow latency","Indexing issues"],
         "a":1, "hint":"fair lending bias acceptance tests schema reliability safety"},
    ]

    with st.form("quiz_form"):
        answers = {}
        for i, item in enumerate(QUIZ, 1):
            st.markdown(f"**Q{i}. {item['q']}**")
            if item["type"] in ("mcq","scenario"):
                answers[i] = st.radio("", item["opts"], key=f"q{i}", label_visibility="collapsed")
            else:
                answers[i] = st.text_input("Your answer", key=f"q{i}", placeholder="Type a constraint…")
            st.divider()
        grade = st.form_submit_button("Grade me")

    if grade:
        total = 0
        for i, item in enumerate(QUIZ, 1):
            hint = item.get("hint","")
            expl_chunks = retrieve_explainer(hint, vec, X, paras, k=2) if doc_loaded else []
            expl = expl_chunks[0] if expl_chunks else "See the lesson sections on acceptance tests, guardrails, and foundations."

            if item["type"] in ("mcq","scenario"):
                picked = item["opts"].index(answers[i]) if answers[i] in item["opts"] else -1
                correct = (picked == item["a"])
                total += int(correct)
                st.markdown(f"<span class='chip {'good' if correct else 'dang'}'>Q{i}</span> "
                            f"{'✅ Correct!' if correct else '❌ Not quite.'}", unsafe_allow_html=True)
                st.info(f"Why: {expl}")
            else:
                ans = answers[i].lower().strip()
                correct = any(r.lower() in ans for r in item["rubric"])
                total += int(correct)
                st.markdown(f"<span class='chip {'good' if correct else 'dang'}'>Q{i}</span> "
                            f"{'✅ Good constraint.' if correct else '❌ Try again.'}", unsafe_allow_html=True)
                st.info(f"Why: {expl}")

        st.success(f"Score: {total}/{len(QUIZ)}")
